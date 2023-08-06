from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from Spitzer.config import config
from Spitzer.result import result

def export(path):
    hosts = result.get_hosts()
    ips = list(hosts.keys())
    ips = sorted(ips, key=lambda x:tuple(map(int, x.split('.'))))
    ports = _get_ports(hosts)


    document = Document(config.get_path() + 'Netwerkservices.docx')
    table = document.add_table(rows=len(ports)+3, cols=len(ips)+1)
    table.style = 'TableGrid'
    table.autofit = True

    #merge header cells
    table.rows[0].cells[0].merge(table.rows[1].cells[0])
    table.rows[0].cells[0].merge(table.rows[2].cells[0])

    for cell in table.rows[0].cells[2:]:
        table.rows[0].cells[1].merge(cell)

    table.rows[0].cells[0].text = 'TCP-poort'
    table.rows[0].cells[1].text = 'IP-adres'

    #write ports
    for i in range(len(ports)):
        table.rows[i+3].cells[0].text = str(ports[i])

    i = 1
    last_range = ''
    for ip in ips:
        p = ip.split('.')[3]
        r = ip.split('.')[0] + '.' + ip.split('.')[1] + '.' + ip.split('.')[2] + '.'
        if last_range != r:
            last_range = r
            table.rows[1].cells[i].text = r  + 'xxx'
        else:
            table.rows[1].cells[i-1].merge(table.rows[1].cells[i])

        table.rows[2].cells[i].text = p
        #change text direction
        tcpr = table.rows[2].cells[i]._tc.get_or_add_tcPr()
        text_direction = OxmlElement('w:textDirection')
        text_direction.set(qn('w:val'), 'tbRl')  # btLr tbRl
        tcpr.append(text_direction)
        
        for port in hosts[ip]:
            index = ports.index(port)
            table.rows[index+3].cells[i].text = 'X'
        i += 1
    
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.name = 'Noto Sans'

    document.save(path)

        

def _get_ports(hosts):
    result = []
    for _, ports in hosts.items():
        for port in ports:
            if port not in result:
                result.append(port)

    result.sort()
    return result