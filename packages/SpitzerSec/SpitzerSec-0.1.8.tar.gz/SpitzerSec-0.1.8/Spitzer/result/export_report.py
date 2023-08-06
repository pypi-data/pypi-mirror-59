from docx import Document
from docx.shared import Pt
from datetime import date

from Spitzer.print import print_error
from Spitzer.config.config import get_path
from Spitzer.result import result

def export(path):
    #TODO  sort for big range

    hosts = result.get_hosts()
    #sort ips
    ips = list(hosts.keys())
    ips = sorted(ips, key=lambda x:tuple(map(int, x.split('.'))))

    #create doc
    document = Document()

    #create standard first page
    document.add_picture(get_path() + 'logo.png')

    paragraph(document, 'VULNERABILITY SCAN REPORT', 'Montserrat', 11, True)
    paragraph(document, 'This is our report of the vulnerability scan conducted on your external IP addresses on '+str(date.today())+'.\
        \nWe’ve scanned the following addresses:', 'Noto Sans', 10)
    p = document.add_paragraph('', style='List Bullet')
    p.add_run().font.name = 'Noto Sans'
    p.add_run().font.size = Pt(9)

    paragraph(document, '\nScans originated from:', 'Noto Sans', 10)

    table = document.add_table(rows=1, cols=3)
    table.style = 'TableGrid'
    table.autofit = True

    paragraph(document, 'Our findings were as follows:', 'Noto Sans', 10)#TODO set font
    p = document.add_paragraph('', style='List Bullet')
    p.add_run().font.name = 'Noto Sans'
    p.add_run().font.size = Pt(9)

    #create table
    table = document.add_table(rows=1, cols=3)
    table.style = 'TableGrid'
    table.autofit = True

    #set header cells
    table.rows[0].cells[0].text = 'Ip Adress'
    table.rows[0].cells[1].text = 'Port'
    table.rows[0].cells[2].text = 'Remarks/Issues'

    #i = 1
    #loop over hosts and write versions
    for host, value in hosts.items():
        #write ip
        row = table.add_row()
        row.cells[0].text = host

        first = True
        #loop over ports
        for port, data in value.items():
            
            #check if first port
            current_row = None
            if not first:
                #add row and merge ip
                current_row = table.add_row()
                row.cells[0].merge(current_row.cells[0])
                #write port number
               
            else:
                first = False
                #write port number
                current_row = row
            
            current_row.cells[1].text = str(port) + '/tcp'

            #write version

            #if everything is unknown
            if data[0] == '' and data[1] == '':
                current_row.cells[2].text = 'Port is open, but service in unknown.'

            #if service is known but not the type
            elif data[0] != '' and data[1] == '':
                current_row.cells[2].text = 'A ' + data[0] + ' service is online.'
            
            #if type is known
            elif data[1] != '':
                current_row.cells[2].text = data[1] + ' ' + data[2] + ' is running.'

            else:
                #uuhhhmm... wtf?!
                print_error('Something went terrebly wrong......')

    #change font of table
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for p in paragraphs:
                for run in p.runs:
                    font = run.font
                    font.name = 'Noto Sans'

    #save document
    document.save(path)



def paragraph(document, text, font, size, bold=False):
    p = document.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.name = font
    run.font.size = Pt(size)
